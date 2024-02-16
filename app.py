
import cv2
import numpy as np
from docx import Document
import sys
import os
import json
import re
import traceback
import pythoncom
import fitz
import base64
import uuid

# Load the pre-trained face detection model
face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')

def extract_images_with_faces_from_pdf(pdf_path, image_folder):
    pdf_document = fitz.open(pdf_path)
    
    # Create the directory if it doesn't exist
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    
    images_with_faces = []
    
    for page_number in range(pdf_document.page_count):
        page = pdf_document.load_page(page_number)
        images = page.get_images(full=True)
        
        for img_index, img_info in enumerate(images):
            xref = img_info[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            
            image = cv2.imdecode(np.frombuffer(image_bytes, np.uint8), cv2.IMREAD_COLOR)
            
            # Convert the image to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Detect faces in the image
            faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5, minSize=(30, 30))
            
            # If faces are detected, save the image
            if len(faces) > 0:
                image_filename = f"page{page_number + 1}_image{img_index + 1}.png"
                image_path = os.path.join(image_folder, image_filename)
                cv2.imwrite(image_path, image)
                images_with_faces.append({"filename": image_filename})
    
    pdf_document.close()
    return images_with_faces

def extract_candidate_name_from_cv(pdf_path):
    try:
        pdf_document = fitz.open(pdf_path)
        candidate_name = ""
        name_patterns = [
            r"Name:\s*(.+)",
            r"Full\s*Name:\s*(.+)",
            r"^([A-Z][a-z]*|[A-Z]+)(?:\s[A-Z][a-z]*)*$",  # Updated regex pattern
        ]
        for page_number in range(len(pdf_document)):
            page = pdf_document.load_page(page_number)
            raw_text = page.get_text('dict', sort=False)
            for block in raw_text['blocks']:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = (span['text']).strip()
                        # for pattern in name_patterns:
                        match = re.search(r"^([A-Z][a-z]*\s?)+$", text)
                        # r"^([A-Z][a-z]*|[A-Z]+\s[A-Z]+)$"
                        # print(f"{match} ~ {span}")
                        if match:
                            candidate_name = match.group(0).strip()
                            break
                    if candidate_name:
                        break
                if candidate_name:
                    break
            if candidate_name:
                break

            # if candidate_name:
            #     break

        pdf_document.close()
        return candidate_name
    except Exception as e:
        # Log the error and return None
        traceback.print_exc()
        return None

def extract_images_from_pdf(pdf_path, image_folder):
    try:
        extracted_images = []
        pdf_document = fitz.open(pdf_path)

        for page_number in range(len(pdf_document)):
            page = pdf_document.load_page(page_number)
            images = page.get_images(full=True)

            for img_index, img_info in enumerate(images):
                xref = img_info[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                # print(image_bytes)
                image_filename = f"{uuid.uuid4()}_image.png"
                image_path = os.path.join(image_folder, image_filename)
                
                with open(image_path, "wb") as image_file:
                    image_file.write(image_bytes)
                
                # Convert image bytes to base64 encoded string
                image_base64 = base64.b64encode(image_bytes).decode('utf-8')

                extracted_images.append({"filename": image_filename})

        pdf_document.close()
        return extracted_images
    except Exception as e:
        # Log the error and return an empty list
        traceback.print_exc()
        return []

# Function to extract text from .docx file
def extract_name_from_docx(docx_path):
    try:
        candidate_name = ""
        doc = Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            # text += paragraph.text + "\n"
            text = paragraph.text.strip()
            # for pattern in name_patterns:
            match = re.search(r"^([A-Z][a-z]*\s?)+$", text)
            if match:
                candidate_name = match.group(0).strip()
                break
        return candidate_name
    except Exception as e:
        # Log the error and return None
        traceback.print_exc()
        return None

# Function to extract images from .docx file
def extract_images_from_docx(docx_path, image_folder):
    images = []
    doc = Document(docx_path)
    for rel_id in doc.part.rels:
        rel = doc.part.rels[rel_id]
        print(rel.reltype)
        if "image" in rel.reltype:
            print(rel.target_part)
            # Get image binary data
            image_part = rel.target_part
            image_bytes = image_part.blob

            # if not image_exists(image_bytes, image_folder):
            # Convert image bytes to base64 encoded string
            # image_base64 = base64.b64encode(image_bytes).decode('utf-8')
            image = cv2.imdecode(np.frombuffer(image_bytes, np.uint8), cv2.IMREAD_COLOR)
            
            # Convert the image to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Detect faces in the image
            faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5, minSize=(30, 30))
            
            # If faces are detected, save the image
            if len(faces) > 0:
                image_filename = f"image_{uuid.uuid4()}.png"
                image_path = os.path.join(image_folder, image_filename)
                cv2.imwrite(image_path, image)
                images.append({"filename": image_filename})
    return images

def main():
    try:
        pythoncom.CoInitialize()
        output_json = {"status": "error", "message": "", "data": {}}
        if len(sys.argv) != 2:
            output_json["message"] = "File not found"
            return json.dumps(output_json, indent=4)

        file_path = sys.argv[1]

        if not os.path.exists(file_path):
            output_json["message"] = "File does not exist"
            return json.dumps(output_json, indent=4)

        if file_path.lower().endswith('.pdf'):
            candidate_name = extract_candidate_name_from_cv(file_path)
            extracted_images = extract_images_with_faces_from_pdf(file_path, 'images')
        elif file_path.lower().endswith('.docx') and file_path.lower().endswith('.doc'):
            candidate_name = extract_name_from_docx(file_path)
            extracted_images = extract_images_from_docx(file_path, 'images')
        else:
            output_json["message"] = "File is not a PDF"
            return json.dumps(output_json, indent=4)
        
        output_json["status"] = "success"
        output_json["message"] = "File processed successfully"
        output_json["data"]["candidate_name"] = candidate_name
        output_json["data"]["extracted_images"] = extracted_images
        return json.dumps(output_json, indent=4)
    except Exception as e:
        traceback.print_exc()
        return json.dumps(output_json, indent=4)
    finally:
        pythoncom.CoUninitialize()
    
if __name__ == "__main__":
    print(main())
    # main()